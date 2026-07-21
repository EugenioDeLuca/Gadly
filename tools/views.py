from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404, resolve_url
from django.http import JsonResponse, HttpResponse, Http404
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth import views as auth_views
from django.contrib.auth.models import User
from django.core.mail import send_mail, EmailMultiAlternatives
from django.conf import settings
from django.urls import reverse
from django.template.loader import render_to_string
from django.templatetags.static import static
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET, require_http_methods
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.utils import timezone
from django.utils.translation import gettext_lazy as _
from django.utils.translation import gettext
from django.utils.translation import override, get_language_from_request
import asyncio
import json
import secrets
import logging
from types import SimpleNamespace
from email.utils import parseaddr
import httpx

from .auth_throttle import (
    SCOPE_PASSWORD_RESET,
    SCOPE_RESEND_VERIFY,
    allow_email_flow,
)
from .forms import (
    UserRegistrationForm,
    ProfileUpdateForm,
    AvatarUploadForm,
    DroseQuoteRequestForm,
    DroseQuoteOfferForm,
)
from .auth_brand import drose_email_brand_preview, email_brand_context, resolve_logout_redirect
from .drose_quotes import (
    mark_quote_viewed_by_staff,
    notify_staff_quote_response,
    process_quote_submission,
    send_quote_offer_to_customer,
)
from .interview_banks import get_interview_simulator_payload
from .models import UserProfile, EmailVerificationToken, DroseQuoteRequest, DroseWorkItem
from .drose_works_media import (
    validate_work_upload,
    validate_work_slot,
    work_items_by_slot,
    save_work_item,
    remove_work_item,
    slot_range,
    attach_gallery_thumb_urls,
)
from PyPDF2 import PdfReader, PdfWriter
from googletrans import Translator  # Importa googletrans per la traduzione
import hashlib
import os
import re
import tempfile
import time
import xml.etree.ElementTree as ET
import zipfile
from urllib.request import urlopen, Request
from urllib.error import URLError, HTTPError
from urllib.parse import quote, urlparse, urlunparse
import socket
import json

logger = logging.getLogger(__name__)


def _request_language(request):
    lang = getattr(request, "LANGUAGE_CODE", None)
    if not lang:
        lang = get_language_from_request(request, check_path=True)
    return lang or settings.LANGUAGE_CODE


class LocalizedPasswordResetView(auth_views.PasswordResetView):
    """Reset email in UI language, with Drose branding when requested from /drose/*."""

    def form_valid(self, form):
        email = form.cleaned_data.get("email") or ""
        allowed, message = allow_email_flow(
            self.request,
            scope=SCOPE_PASSWORD_RESET,
            email=email,
        )
        if not allowed:
            # Neutral message: does not reveal whether the address exists.
            form.add_error(None, message)
            return self.form_invalid(form)
        self.extra_email_context = email_brand_context(self.request)
        with override(_request_language(self.request)):
            return super().form_valid(form)


def _send_via_sendgrid_api(from_email, to_email, subject, plain_message, html_message):
    api_key = (os.environ.get("SENDGRID_API_KEY") or "").strip()
    if not api_key:
        return False
    name, email_addr = parseaddr(from_email or "")
    sender_email = (email_addr or "").strip() or "contact@gadly.it"
    sender_name = (name or "").strip() or "Gadly Support"
    payload = {
        "personalizations": [{"to": [{"email": to_email}]}],
        "from": {"email": sender_email, "name": sender_name},
        "subject": subject,
        "content": [
            {"type": "text/plain", "value": plain_message},
            {"type": "text/html", "value": html_message},
        ],
    }
    try:
        timeout_seconds = int((os.environ.get("EMAIL_TIMEOUT") or "20").strip() or "20")
    except Exception:
        timeout_seconds = 20
    try:
        resp = httpx.post(
            "https://api.sendgrid.com/v3/mail/send",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=payload,
            timeout=timeout_seconds,
        )
        if 200 <= resp.status_code < 300:
            return True
        logger.error("SendGrid API send failed status=%s body=%s", resp.status_code, resp.text[:500])
    except Exception:
        logger.exception("SendGrid API send failed unexpectedly")
    return False


def _send_verification_email(request, user, token, *, email_change=False, to_email=None):
    verify_url = request.build_absolute_uri(
        reverse("verify_email_short", kwargs={"token": token})
    )
    brand = email_brand_context(request)
    is_drose = brand["is_drose_brand"]
    with override(_request_language(request)):
        if email_change:
            if is_drose:
                subject = gettext("Confirm your new email - Drose Drone Services")
                plain_message = gettext(
                    "Hi %(username)s,\n\n"
                    "You recently changed the email address on your Gadly account "
                    "(used for Gadly tools and Drose Drone Services). "
                    "Please confirm this new address by opening the following link:\n\n"
                    "%(url)s\n\n"
                    "Confirming this address keeps your account secure and lets you use it for password recovery.\n\n"
                    "For your security, this verification link is valid for a limited time and can be used only once.\n\n"
                    "If you did not change your email, you can ignore this message.\n\n"
                    "Best regards,\n"
                    "Drose Drone Services"
                ) % {"username": user.username, "url": verify_url}
            else:
                subject = gettext("Confirm your new email - Gadly")
                plain_message = gettext(
                    "Hi %(username)s,\n\n"
                    "You recently changed the email address on your Gadly account. "
                    "Please confirm this new address by opening the following link:\n\n"
                    "%(url)s\n\n"
                    "Confirming this address keeps your account secure and lets you use it for password recovery.\n\n"
                    "For your security, this verification link is valid for a limited time and can be used only once.\n\n"
                    "If you did not change your email on Gadly, you can ignore this email.\n\n"
                    "Best regards,\n"
                    "Gadly Support"
                ) % {"username": user.username, "url": verify_url}
        elif is_drose:
            subject = gettext("Verify your email - Drose Drone Services")
            plain_message = gettext(
                "Hi %(username)s,\n\n"
                "Thanks for signing up from Drose Drone Services. "
                "Please verify your email address by clicking this link:\n\n"
                "%(url)s\n\n"
                "Your Gadly account works for both Gadly tools and Drose pages.\n\n"
                "For your security, this verification link is valid for a limited time and can be used only once.\n\n"
                "If you didn't create an account, ignore this email.\n\n"
                "Best regards,\n"
                "Drose Drone Services"
            ) % {"username": user.username, "url": verify_url}
        else:
            subject = gettext("Verify your email - Gadly")
            plain_message = gettext(
                "Hi %(username)s,\n\n"
                "Please verify your email address by clicking this link:\n\n%(url)s\n\n"
                "This confirmation helps keep your account secure, completes your account setup, and enables account recovery.\n\n"
                "For your security, this verification link is valid for a limited time and can be used only once.\n\n"
                "If you didn't create an account, ignore this email.\n\n"
                "Best regards,\n"
                "Gadly Support"
            ) % {"username": user.username, "url": verify_url}
        html_message = render_to_string(
            "tools/emails/verify_email.html",
            {
                "username": user.username,
                "verify_url": verify_url,
                "email_change": email_change,
                **brand,
            },
            request=request,
        )
    recipient = (to_email or user.email or "").strip()
    if not recipient:
        return None
    from_email = getattr(settings, 'DEFAULT_FROM_EMAIL', 'noreply@gadly.local')
    msg = EmailMultiAlternatives(subject, plain_message, from_email, [recipient])
    msg.attach_alternative(html_message, "text/html")
    smtp_ok = False
    api_ok = False
    reason = "unknown"
    try:
        msg.send(fail_silently=False)
        smtp_ok = True
        reason = "smtp_sent"
        logger.info(
            "EMAIL_FLOW user=%s email=%s user_found=1 smtp_ok=1 api_ok=0 reason=%s",
            user.username,
            user.email,
            reason,
        )
        return verify_url
    except Exception:
        logger.exception("SMTP send failed for verification email user=%s", user.email)
        reason = "smtp_failed"
    if _send_via_sendgrid_api(from_email, recipient, subject, plain_message, html_message):
        api_ok = True
        reason = "api_sent_after_smtp_fail"
        logger.info(
            "EMAIL_FLOW user=%s email=%s user_found=1 smtp_ok=%d api_ok=%d reason=%s",
            user.username,
            user.email,
            1 if smtp_ok else 0,
            1 if api_ok else 0,
            reason,
        )
        return verify_url
    reason = "smtp_and_api_failed"
    logger.error(
        "EMAIL_FLOW user=%s email=%s user_found=1 smtp_ok=%d api_ok=%d reason=%s",
        user.username,
        user.email,
        1 if smtp_ok else 0,
        1 if api_ok else 0,
        reason,
    )
    return None


def _sanitize_home_hidden_tools(data):
    """Normalize hidden-tool entries from client JSON or profile storage."""
    if not isinstance(data, list):
        return []
    seen = set()
    out = []
    for item in data:
        if not isinstance(item, dict):
            continue
        url = (item.get("url") or "").strip()
        if not url.startswith("/"):
            continue
        if len(url) > 1 and url.endswith("/"):
            url = url[:-1]
        if not url or url in seen:
            continue
        seen.add(url)
        out.append({
            "url": url,
            "name": str(item.get("name") or url)[:200],
            "categoryId": str(item.get("categoryId") or "")[:100],
            "categoryName": str(item.get("categoryName") or "")[:200],
        })
    return out


# -------------------------
# Normal page views
# -------------------------
def home(request):
    return render(request, "tools/home.html")


def drose(request):
    from .drose_service_thumbs import drose_service_photo_uris

    return render(
        request,
        "tools/drose.html",
        {"drose_service_photos": drose_service_photo_uris()},
    )


def drose_quote_request(request):
    quote_form = DroseQuoteRequestForm()
    quote_sent = False

    if request.method == "POST":
        quote_form = DroseQuoteRequestForm(request.POST)
        if quote_form.is_valid():
            lang = _request_language(request)
            try:
                process_quote_submission(request, quote_form, language=lang)
                quote_sent = True
                quote_form = DroseQuoteRequestForm()
            except Exception:
                logger.exception("Drose quote form send failed")
                quote_form.add_error(
                    None,
                    gettext("We could not send your request right now. Please try again in a few minutes."),
                )

    return render(
        request,
        'tools/drose_quote_request.html',
        {
            'quote_form': quote_form,
            'quote_sent': quote_sent,
        },
    )


@staff_member_required
def drose_quote_staff_list(request):
    quotes = DroseQuoteRequest.objects.filter(deleted_at__isnull=True)[:200]
    return render(request, 'tools/drose_quote_staff_list.html', {'quotes': quotes})


@staff_member_required
def drose_quote_staff_trash(request):
    quotes = DroseQuoteRequest.objects.filter(deleted_at__isnull=False)[:200]
    return render(
        request,
        'tools/drose_quote_staff_trash.html',
        {
            'quotes': quotes,
            'trash_retention_days': getattr(settings, 'DROSE_QUOTE_TRASH_RETENTION_DAYS', 60),
        },
    )


@staff_member_required
@require_http_methods(["POST"])
def drose_quote_staff_delete(request, pk):
    quote = get_object_or_404(DroseQuoteRequest, pk=pk, deleted_at__isnull=True)
    quote.deleted_at = timezone.now()
    quote.save(update_fields=["deleted_at", "updated_at"])
    return redirect('drose_quote_staff_list')


@staff_member_required
@require_http_methods(["POST"])
def drose_quote_staff_restore(request, pk):
    quote = get_object_or_404(DroseQuoteRequest, pk=pk, deleted_at__isnull=False)
    quote.deleted_at = None
    quote.save(update_fields=["deleted_at", "updated_at"])
    return redirect(f"{reverse('drose_quote_staff_list')}?restored={quote.pk}")


@staff_member_required
@require_http_methods(["POST"])
def drose_quote_staff_purge(request, pk):
    quote = get_object_or_404(DroseQuoteRequest, pk=pk, deleted_at__isnull=False)
    quote.delete()
    return redirect('drose_quote_staff_trash')


@staff_member_required
@require_http_methods(["GET", "POST"])
def drose_quote_staff_detail(request, pk):
    quote = get_object_or_404(DroseQuoteRequest, pk=pk, deleted_at__isnull=True)
    if request.method == "GET":
        mark_quote_viewed_by_staff(quote)
    offer_sent = False
    offer_error = ""

    initial_deliverables = "\n".join(quote.offer_deliverables) if quote.offer_deliverables else quote.deliverables_label
    initial = {
        "valid_until": quote.valid_until,
        "offer_scope": quote.offer_scope or quote.project_details,
        "offer_deliverables_text": initial_deliverables,
        "price_subtotal": quote.price_subtotal,
        "price_total": quote.price_total,
        "offer_timeline": quote.offer_timeline or quote.preferred_timeline,
        "offer_notes": quote.offer_notes,
    }
    offer_form = DroseQuoteOfferForm(initial=initial)

    if request.method == "POST":
        offer_form = DroseQuoteOfferForm(request.POST)
        if offer_form.is_valid():
            quote.valid_until = offer_form.cleaned_data["valid_until"]
            quote.offer_scope = offer_form.cleaned_data["offer_scope"]
            quote.offer_deliverables = offer_form.deliverables_list()
            quote.price_subtotal = offer_form.cleaned_data["price_subtotal"]
            quote.price_total = offer_form.cleaned_data["price_total"]
            quote.offer_timeline = offer_form.cleaned_data["offer_timeline"]
            quote.offer_notes = offer_form.cleaned_data.get("offer_notes") or ""
            quote.save()
            try:
                send_quote_offer_to_customer(request, quote)
                offer_sent = True
            except Exception:
                logger.exception("Drose quote offer send failed for %s", quote.reference)
                offer_error = gettext("We could not send the quote email. Check mail settings and try again.")

    return render(
        request,
        'tools/drose_quote_staff_detail.html',
        {
            'quote': quote,
            'offer_form': offer_form,
            'offer_sent': offer_sent,
            'offer_error': offer_error,
        },
    )


@require_http_methods(["GET", "POST"])
def drose_quote_response(request, token):
    quote = get_object_or_404(DroseQuoteRequest, response_token=token)
    intent = (request.GET.get("intent") or request.POST.get("intent") or "").strip().lower()
    completed = False
    completed_action = ""

    if request.method == "POST" and quote.can_respond_to_offer:
        action = (request.POST.get("action") or "").strip().lower()
        if action == "accept":
            quote.status = DroseQuoteRequest.STATUS_ACCEPTED
            quote.responded_at = timezone.now()
            quote.staff_viewed_at = None
            quote.save(update_fields=["status", "responded_at", "staff_viewed_at", "updated_at"])
            notify_staff_quote_response(request, quote)
            completed = True
            completed_action = "accepted"
        elif action == "decline":
            quote.status = DroseQuoteRequest.STATUS_DECLINED
            quote.responded_at = timezone.now()
            quote.staff_viewed_at = None
            quote.save(update_fields=["status", "responded_at", "staff_viewed_at", "updated_at"])
            notify_staff_quote_response(request, quote)
            completed = True
            completed_action = "declined"

    with override(quote.language):
        return render(
            request,
            'tools/drose_quote_response.html',
            {
                'quote': quote,
                'intent': intent if intent in ("accept", "decline") else "",
                'completed': completed,
                'completed_action': completed_action,
                'can_respond': quote.can_respond_to_offer,
            },
        )


def drose_faq(request):
    return render(request, 'tools/drose_faq.html')


def drose_about(request):
    return render(request, 'tools/drose_about.html')


def drose_works(request):
    raw_type = (request.GET.get("type") or "").strip().lower()
    works_media_type = raw_type if raw_type in (
        DroseWorkItem.MEDIA_PHOTO,
        DroseWorkItem.MEDIA_VIDEO,
    ) else ""
    # Mobile intermediate step: ?choose=1 shows Photos/Videos picker without loading grids.
    show_works_chooser = request.GET.get("choose") == "1" and not works_media_type

    photo_items = work_items_by_slot(DroseWorkItem.MEDIA_PHOTO)
    video_items = work_items_by_slot(DroseWorkItem.MEDIA_VIDEO)
    has_any_media = bool(photo_items or video_items)

    if show_works_chooser:
        photo_slots = []
        video_slots = []
        has_work_media = has_any_media
    else:
        attach_gallery_thumb_urls(photo_items)
        show_photos = works_media_type in ("", DroseWorkItem.MEDIA_PHOTO)
        show_videos = works_media_type in ("", DroseWorkItem.MEDIA_VIDEO)
        photo_slots = (
            [(n, photo_items.get(n)) for n in slot_range(DroseWorkItem.MEDIA_PHOTO)]
            if show_photos
            else []
        )
        video_slots = (
            [(n, video_items.get(n)) for n in slot_range(DroseWorkItem.MEDIA_VIDEO)]
            if show_videos
            else []
        )
        if works_media_type == DroseWorkItem.MEDIA_PHOTO:
            has_work_media = bool(photo_items)
        elif works_media_type == DroseWorkItem.MEDIA_VIDEO:
            has_work_media = bool(video_items)
        else:
            has_work_media = has_any_media

    # data-URI già in pagina: niente preload (evita doppio fetch / race)
    preload_urls = []
    return render(
        request,
        'tools/drose_works.html',
        {
            'photo_slots': photo_slots,
            'video_slots': video_slots,
            'can_manage_works': request.user.is_authenticated and request.user.is_staff,
            'has_work_media': has_work_media,
            'work_preload_urls': preload_urls,
            'works_media_type': works_media_type,
            'show_works_chooser': show_works_chooser,
        },
    )


@staff_member_required
def drose_works_manage(request):
    from django.core.exceptions import ValidationError

    upload_error = ""
    upload_success = False
    removed = False

    if request.method == "POST":
        action = request.POST.get("action")
        media_type = request.POST.get("media_type")
        slot = request.POST.get("slot")
        if action == "remove":
            try:
                media_type, slot_num = validate_work_slot(media_type, slot)
                if remove_work_item(media_type, slot_num):
                    removed = True
                else:
                    upload_error = gettext("Could not remove this item.")
            except ValidationError as exc:
                upload_error = exc.messages[0] if exc.messages else str(exc)
        elif action == "upload":
            uploaded = request.FILES.get("file")
            try:
                media_type, slot_num = validate_work_upload(media_type, slot, uploaded)
                save_work_item(media_type, slot_num, uploaded)
                upload_success = True
            except ValidationError as exc:
                upload_error = exc.messages[0] if exc.messages else str(exc)
            except Exception:
                logger.exception("Drose work upload failed")
                upload_error = gettext("Upload failed. Please try again.")

    photo_items = work_items_by_slot(DroseWorkItem.MEDIA_PHOTO)
    video_items = work_items_by_slot(DroseWorkItem.MEDIA_VIDEO)
    attach_gallery_thumb_urls(photo_items)
    photo_slots = [(n, photo_items.get(n)) for n in slot_range(DroseWorkItem.MEDIA_PHOTO)]
    video_slots = [(n, video_items.get(n)) for n in slot_range(DroseWorkItem.MEDIA_VIDEO)]
    preload_urls = []
    return render(
        request,
        'tools/drose_works_manage.html',
        {
            'photo_slots': photo_slots,
            'video_slots': video_slots,
            'upload_error': upload_error,
            'upload_success': upload_success,
            'removed': removed,
            'work_preload_urls': preload_urls,
        },
    )


def drose_privacy(request):
    return render(request, 'tools/drose_privacy.html')


def drose_contact(request):
    return render(request, 'tools/drose_contact.html')


def _ensure_debug_email_preview():
    if not settings.DEBUG:
        raise Http404


@require_GET
def drose_email_preview_index(request):
    """Index of Drose transactional email previews (DEBUG only)."""
    _ensure_debug_email_preview()
    return render(request, 'tools/drose_email_preview_index.html')


@require_GET
def drose_email_preview_verify(request):
    """Render verify-email HTML as sent from Drose (DEBUG only)."""
    _ensure_debug_email_preview()
    lang = (request.GET.get('lang') or get_language_from_request(request) or 'en')[:2]
    if lang not in ('en', 'it'):
        lang = 'en'
    email_change = request.GET.get('change') == '1'
    with override(lang):
        html = render_to_string(
            'tools/emails/verify_email.html',
            {
                'username': 'mario',
                'verify_url': request.build_absolute_uri('/v/example-token-preview/'),
                'email_change': email_change,
                **drose_email_brand_preview(request),
            },
            request=request,
        )
    return HttpResponse(html)


@require_GET
def drose_email_preview_password_reset(request):
    """Render password-reset HTML as sent from Drose (DEBUG only)."""
    _ensure_debug_email_preview()
    lang = (request.GET.get('lang') or get_language_from_request(request) or 'en')[:2]
    if lang not in ('en', 'it'):
        lang = 'en'
    user = SimpleNamespace(get_username=lambda: 'mario')
    with override(lang):
        html = render_to_string(
            'registration/password_reset_email_html.html',
            {
                'user': user,
                'protocol': 'https' if request.is_secure() else 'http',
                'domain': request.get_host(),
                'uid': 'preview',
                'token': 'preview-token',
                'site_name': 'Gadly',
                **drose_email_brand_preview(request),
            },
            request=request,
        )
    return HttpResponse(html)


@require_GET
def drose_email_preview_quote_offer(request):
    """Render quote-offer HTML template with Drose branding (DEBUG only)."""
    _ensure_debug_email_preview()
    lang = (request.GET.get('lang') or get_language_from_request(request) or 'en')[:2]
    if lang not in ('en', 'it'):
        lang = 'en'
    with override(lang):
        html = render_to_string(
            'tools/emails/drose_quote_offer.html',
            {
                'recipient_name': 'Mario Rossi',
                'quote_reference': 'DS-2026-041',
                'valid_until': '31/07/2026',
                'location': 'Milano',
                'service_label': gettext('Inspection'),
                'project_scope': gettext('Facade and roof visual inspection with 4K coverage.'),
                'deliverables': [
                    gettext('12 edited aerial photos'),
                    gettext('1 short recap video (45-60 seconds)'),
                    gettext('Inspection summary PDF'),
                ],
                'price_subtotal': '680 EUR',
                'price_total': '680 EUR',
                'timeline': gettext('Delivery in 3 business days after flight.'),
                'notes': gettext('Flight date subject to weather and local airspace clearance.'),
                'accept_url': request.build_absolute_uri(reverse('drose_quote_response', kwargs={'token': 'preview-token'})) + '?intent=accept',
                'decline_url': request.build_absolute_uri(reverse('drose_quote_response', kwargs={'token': 'preview-token'})) + '?intent=decline',
                **drose_email_brand_preview(request),
            },
            request=request,
        )
    return HttpResponse(html)


@require_GET
def drose_email_preview_quote_received(request):
    """Render quote-request-received HTML template with Drose branding (DEBUG only)."""
    _ensure_debug_email_preview()
    lang = (request.GET.get('lang') or get_language_from_request(request) or 'en')[:2]
    if lang not in ('en', 'it'):
        lang = 'en'
    with override(lang):
        html = render_to_string(
            'tools/emails/drose_quote_request_received.html',
            {
                'recipient_name': 'Mario Rossi',
                'quote_reference': 'DS-2026-0041',
                **drose_email_brand_preview(request),
            },
            request=request,
        )
    return HttpResponse(html)


def hidden_tools(request):
    return render(request, 'tools/hidden_tools.html')


@login_required
@require_http_methods(["GET", "POST"])
def api_home_hidden_tools(request):
    """Read/write hidden home tools for cross-device sync (logged-in users)."""
    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    if request.method == "GET":
        stored = profile.home_hidden_tools if isinstance(profile.home_hidden_tools, list) else []
        return JsonResponse({"tools": _sanitize_home_hidden_tools(stored)})
    try:
        body = json.loads(request.body.decode() or "{}")
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({"error": "invalid_json"}, status=400)
    tools = _sanitize_home_hidden_tools(body.get("tools", []))
    profile.home_hidden_tools = tools
    profile.save(update_fields=["home_hidden_tools"])
    return JsonResponse({"ok": True, "tools": tools})


def faq(request):
    return render(request, 'tools/faq.html')


def privacy(request):
    return render(request, 'tools/privacy.html')


def about(request):
    return render(request, 'tools/about.html')


def contact(request):
    return render(request, 'tools/contact.html')


class DroseAwareLogoutView(auth_views.LogoutView):
    """Logout with redirect to /drose/ when the session started from Drose."""

    http_method_names = ["post", "options"]
    next_page = None

    def dispatch(self, request, *args, **kwargs):
        self._drose_logout_target = resolve_logout_redirect(request)
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        target = getattr(self, "_drose_logout_target", "")
        if target:
            return target
        redirect_url = self.get_redirect_url()
        if redirect_url:
            return redirect_url
        if settings.LOGOUT_REDIRECT_URL:
            return resolve_url(settings.LOGOUT_REDIRECT_URL)
        return "/"


def accounts_home(request):
    """Redirect /accounts/ to admin (staff), account profile (user), or login."""
    if request.user.is_authenticated and request.user.is_staff:
        from django.urls import reverse
        return redirect(reverse('admin:auth_user_changelist'))
    if request.user.is_authenticated:
        return redirect('account_profile')
    return redirect('login')


def _profile_form_wants_ajax(request):
    return request.headers.get("X-Gadly-Profile-Ajax") == "1"


def _profile_form_ajax_success(request, profile, message=None):
    user = request.user
    username = (user.username or "").strip()
    return JsonResponse({
        "ok": True,
        "message": message or gettext("Profile updated."),
        "username": username,
        "email": user.email or "",
        "pending_email": profile.pending_email or "",
        "initials": (username[:2].upper() if username else "?"),
    })


def _profile_form_ajax_redirect(url):
    return JsonResponse({"ok": True, "redirect": url})


def _profile_form_ajax_errors(profile_form):
    errors = {}
    for field, field_errors in profile_form.errors.items():
        key = field if field != "__all__" else "_all"
        errors[key] = [str(err) for err in field_errors]
    return JsonResponse({"ok": False, "errors": errors}, status=400)


def _profile_form_ajax_avatar_success(request, profile, message=None):
    from .avatar_inline import avatar_inline_data_uri

    user = request.user
    username = (user.username or "").strip()
    avatar = profile.avatar
    has_avatar = bool(avatar and getattr(avatar, "name", None))
    return JsonResponse({
        "ok": True,
        "message": message or gettext("Profile image saved."),
        "has_avatar": has_avatar,
        "avatar_inline_header": avatar_inline_data_uri(avatar, 56) if has_avatar else "",
        "avatar_inline_account": avatar_inline_data_uri(avatar, 160) if has_avatar else "",
        "initials": (username[:2].upper() if username else "?"),
    })


def _profile_form_ajax_avatar_errors(avatar_form):
    errors = {}
    for field, field_errors in avatar_form.errors.items():
        key = field if field != "__all__" else "_all"
        errors[key] = [str(err) for err in field_errors]
    return JsonResponse({"ok": False, "errors": errors}, status=400)


@login_required
def account_profile(request):
    """Account settings: username, email, avatar, password change."""
    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    profile_form = ProfileUpdateForm(user=request.user)
    avatar_form = AvatarUploadForm(instance=profile)
    msg = None
    avatar_saved = False
    avatar_removed = False
    profile_saved = False
    pending_cancelled = False

    if request.method == 'POST':
        if 'update_profile' in request.POST:
            profile_form = ProfileUpdateForm(request.POST, user=request.user)
            if profile_form.is_valid():
                current_email = (request.user.email or "").strip().lower()
                new_email = profile_form.cleaned_data["email"].strip().lower()
                new_username = profile_form.cleaned_data["username"]
                email_changed = current_email != new_email
                pending_email = (profile.pending_email or "").strip().lower()

                if email_changed:
                    EmailVerificationToken.objects.filter(user=request.user).delete()
                    token = secrets.token_hex(24)
                    EmailVerificationToken.objects.create(user=request.user, token=token)
                    verify_url = _send_verification_email(
                        request,
                        request.user,
                        token,
                        email_change=True,
                        to_email=new_email,
                    )
                    if verify_url:
                        profile.pending_email = new_email
                        profile.save(update_fields=["pending_email"])
                        if new_username != request.user.username:
                            request.user.username = new_username
                            request.user.save(update_fields=["username"])
                        request.session["pending_verify_user_id"] = request.user.id
                        if settings.DEBUG:
                            request.session["pending_verify_url"] = verify_url
                        verify_redirect = reverse("verify_email_sent") + "?email_change=1&sent=1"
                        if _profile_form_wants_ajax(request):
                            return _profile_form_ajax_redirect(verify_redirect)
                        return redirect(verify_redirect)
                    profile_form.add_error(
                        None,
                        gettext(
                            "We could not send the confirmation email. Your email address was not changed."
                        ),
                    )
                elif pending_email and new_email == current_email:
                    profile.pending_email = ""
                    profile.save(update_fields=["pending_email"])
                    request.user.username = new_username
                    request.user.save(update_fields=["username"])
                    if _profile_form_wants_ajax(request):
                        return _profile_form_ajax_success(request, profile)
                    return redirect(
                        reverse("account_profile")
                        + "?profile_saved=1&pending_cancelled=1"
                    )
                elif pending_email:
                    if new_username != request.user.username:
                        request.user.username = new_username
                        request.user.save(update_fields=["username"])
                        if _profile_form_wants_ajax(request):
                            return _profile_form_ajax_success(request, profile)
                        return redirect(
                            reverse("account_profile") + "?profile_saved=1"
                        )
                    profile_form.add_error(
                        "email",
                        gettext(
                            "Confirm the new email address using the link we sent before saving again."
                        ),
                    )
                    if _profile_form_wants_ajax(request):
                        return _profile_form_ajax_errors(profile_form)
                    profile_form = ProfileUpdateForm(user=request.user)
                else:
                    request.user.username = new_username
                    request.user.save(update_fields=["username"])
                    if _profile_form_wants_ajax(request):
                        return _profile_form_ajax_success(request, profile)
                    return redirect(
                        reverse("account_profile") + "?profile_saved=1"
                    )
            if _profile_form_wants_ajax(request) and profile_form.errors:
                return _profile_form_ajax_errors(profile_form)
        elif 'update_avatar' in request.POST:
            wants_avatar_ajax = _profile_form_wants_ajax(request)
            if request.POST.get('avatar_action') == 'remove':
                if profile.avatar:
                    profile.avatar.delete(save=False)
                profile.avatar = None
                profile.save(update_fields=['avatar'])
                if wants_avatar_ajax:
                    return _profile_form_ajax_avatar_success(
                        request,
                        profile,
                        gettext("Profile image removed."),
                    )
                return redirect(reverse('account_profile') + '?avatar_removed=1')
            avatar_form = AvatarUploadForm(request.POST, request.FILES, instance=profile)
            if avatar_form.is_valid():
                media_avatars = os.path.join(settings.MEDIA_ROOT, 'avatars')
                os.makedirs(media_avatars, exist_ok=True)
                avatar_form.save()
                profile.refresh_from_db(fields=['avatar'])
                if request.POST.get('avatar_action') == 'save':
                    if wants_avatar_ajax:
                        return _profile_form_ajax_avatar_success(request, profile)
                    return redirect(reverse('account_profile') + '?avatar_saved=1')
                if wants_avatar_ajax:
                    return _profile_form_ajax_avatar_success(request, profile)
                avatar_form = AvatarUploadForm(instance=profile)
            elif wants_avatar_ajax:
                return _profile_form_ajax_avatar_errors(avatar_form)
    else:
        avatar_saved = request.GET.get('avatar_saved') == '1'
        avatar_removed = request.GET.get('avatar_removed') == '1'
        profile_saved = request.GET.get('profile_saved') == '1'
        pending_cancelled = request.GET.get('pending_cancelled') == '1'

    return render(request, 'tools/account_profile.html', {
        'profile': profile,
        'profile_form': profile_form,
        'avatar_form': avatar_form,
        'msg': msg,
        'avatar_saved': avatar_saved,
        'avatar_removed': avatar_removed,
        'profile_saved': profile_saved,
        'pending_cancelled': pending_cancelled,
    })


@login_required
def account_delete_confirm(request):
    """Confirmation page: enter password to delete account."""
    err = None
    if request.method == 'POST':
        password = request.POST.get('password', '')
        user = authenticate(request, username=request.user.username, password=password)
        if user is not None:
            logout(request)
            user.delete()
            return redirect('home')
        err = gettext('Incorrect password.')
    return render(request, 'tools/account_delete_confirm.html', {'err': err})


@require_http_methods(["GET"])
def check_username(request):
    """API: check if username is available. Returns JSON {available: bool}."""
    username = (request.GET.get('username') or '').strip()
    if not username:
        return JsonResponse({'available': False, 'error': 'empty'})
    if len(username) < 3:
        return JsonResponse({'available': False, 'error': 'too_short'})
    taken = User.objects.filter(username__iexact=username).exists()
    return JsonResponse({'available': not taken})


def register(request):
    if request.user.is_authenticated:
        from tools.auth_brand import resolve_auth_next
        auth_next = resolve_auth_next(request)
        if auth_next:
            return redirect(auth_next)
        return redirect('home')
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            UserProfile.objects.get_or_create(user=user)
            # Send verification email (use hex token - no special chars, avoids copy issues)
            EmailVerificationToken.objects.filter(user=user).delete()
            token = secrets.token_hex(24)
            EmailVerificationToken.objects.create(user=user, token=token)
            verify_url = _send_verification_email(request, user, token)
            request.session['pending_verify_user_id'] = user.id
            # In dev (DEBUG): store link in session so user can click it directly from the page
            if settings.DEBUG and verify_url:
                request.session['pending_verify_url'] = verify_url
            if verify_url:
                return redirect(reverse('verify_email_sent') + '?sent=1')
            return redirect(reverse('verify_email_sent') + '?resent=0')
    else:
        form = UserRegistrationForm()
    return render(request, 'tools/register.html', {'form': form})


@require_GET
def verify_email(request, token):
    """Verify email when user clicks the link."""
    token = (token or '').strip()
    evt = EmailVerificationToken.objects.filter(token=token).first()
    if not evt:
        return render(request, 'tools/verify_email_invalid.html')
    user = evt.user
    profile, _ = UserProfile.objects.get_or_create(user=user)
    pending = (profile.pending_email or "").strip()
    email_change_confirmed = bool(pending)
    confirmed_email = pending
    if pending:
        user.email = pending
        profile.pending_email = ""
        user.save(update_fields=["email"])
        profile.email_verified = True
        profile.save(update_fields=["email_verified", "pending_email"])
    else:
        profile.email_verified = True
        profile.save(update_fields=["email_verified"])
    if request.session.get('pending_verify_user_id') == evt.user.id:
        try:
            del request.session['pending_verify_user_id']
        except Exception:
            pass
    try:
        del request.session['pending_verify_url']
    except Exception:
        pass
    evt.delete()

    if email_change_confirmed:
        if request.user.is_authenticated and request.user.pk == user.pk:
            logout(request)
        login_url = reverse('login') + '?email_changed=1'
        if confirmed_email:
            login_url += '&email=' + quote(confirmed_email)
        return render(request, 'tools/verify_email_done.html', {
            'email_change_confirmed': True,
            'confirmed_email': confirmed_email,
            'login_url': login_url,
        })

    return render(request, 'tools/verify_email_done.html')


def verify_email_sent(request):
    """Page shown after registration: check your email."""
    # In DEBUG mode, show clickable link (email goes to console, link is easier here)
    verify_url = request.session.pop('pending_verify_url', None)
    sent = request.GET.get('sent') == '1'
    resent = request.GET.get('resent') == '1'
    already_verified = request.GET.get('already_verified') == '1'
    resend_failed = (request.GET.get('resent') == '0') and (not already_verified)
    email_change = request.GET.get('email_change') == '1'
    return render(request, 'tools/verify_email_sent.html', {
        'verify_url': verify_url,
        'sent': sent,
        'resent': resent,
        'resend_failed': resend_failed,
        'already_verified': already_verified,
        'email_change': email_change,
    })


def resend_verification_email(request):
    """Resend verification email for logged-in or just-registered users."""
    target_user = request.user if request.user.is_authenticated else None
    if target_user is None:
        pending_user_id = request.session.get('pending_verify_user_id')
        if pending_user_id:
            target_user = User.objects.filter(pk=pending_user_id).first()
    if target_user is None:
        return redirect('login')

    throttle_email = (getattr(target_user, "email", "") or "").strip()
    profile_preview, _ = UserProfile.objects.get_or_create(user=target_user)
    pending_change_preview = (profile_preview.pending_email or "").strip()
    if pending_change_preview:
        throttle_email = pending_change_preview
    allowed, message = allow_email_flow(
        request,
        scope=SCOPE_RESEND_VERIFY,
        email=throttle_email,
    )
    if not allowed:
        messages.error(request, message)
        return redirect(reverse('verify_email_sent') + '?resent=0')

    profile, _ = UserProfile.objects.get_or_create(user=target_user)
    pending_change = (profile.pending_email or "").strip()
    if pending_change:
        token = secrets.token_hex(24)
        new_evt = EmailVerificationToken.objects.create(user=target_user, token=token)
        verify_url = _send_verification_email(
            request,
            target_user,
            token,
            email_change=True,
            to_email=pending_change,
        )
        request.session['pending_verify_user_id'] = target_user.id
        if settings.DEBUG and verify_url:
            request.session['pending_verify_url'] = verify_url
        if verify_url:
            EmailVerificationToken.objects.filter(user=target_user).exclude(pk=new_evt.pk).delete()
            return redirect(reverse('verify_email_sent') + '?email_change=1&resent=1')
        new_evt.delete()
        return redirect(reverse('verify_email_sent') + '?email_change=1&resent=0')
    # Keep profile flag aligned for legacy accounts that are clearly active.
    if (not profile.email_verified) and target_user.last_login is not None:
        profile.email_verified = True
        profile.save(update_fields=["email_verified"])
    if profile.email_verified:
        return redirect(reverse('verify_email_sent') + '?already_verified=1')
    token = secrets.token_hex(24)
    new_evt = EmailVerificationToken.objects.create(user=target_user, token=token)
    verify_url = _send_verification_email(request, target_user, token)
    request.session['pending_verify_user_id'] = target_user.id
    if settings.DEBUG and verify_url:
        request.session['pending_verify_url'] = verify_url
    if verify_url:
        EmailVerificationToken.objects.filter(user=target_user).exclude(pk=new_evt.pk).delete()
        return redirect(reverse('verify_email_sent') + '?resent=1')
    new_evt.delete()
    return redirect(reverse('verify_email_sent') + '?resent=0')


@require_http_methods(["POST"])
def resend_verification_email_public(request):
    email = (request.POST.get("email") or "").strip()
    allowed, message = allow_email_flow(
        request,
        scope=SCOPE_RESEND_VERIFY,
        email=email or "missing",
    )
    if not allowed:
        messages.error(request, message)
        return redirect(reverse('request_verification_email_page'))

    if email:
        user = User.objects.filter(email__iexact=email).first()
        if user:
            profile, _ = UserProfile.objects.get_or_create(user=user)
            # Defensive alignment for legacy accounts that already logged in.
            if (not profile.email_verified) and user.last_login is not None:
                profile.email_verified = True
                profile.save(update_fields=["email_verified"])
            if profile.email_verified:
                return redirect(reverse('verify_email_sent') + '?already_verified=1')
            token = secrets.token_hex(24)
            new_evt = EmailVerificationToken.objects.create(user=user, token=token)
            verify_url = _send_verification_email(request, user, token)
            if settings.DEBUG and verify_url:
                request.session['pending_verify_url'] = verify_url
            if verify_url:
                EmailVerificationToken.objects.filter(user=user).exclude(pk=new_evt.pk).delete()
                return redirect(reverse('verify_email_sent') + '?resent=1')
            new_evt.delete()
            return redirect(reverse('verify_email_sent') + '?resent=0')
        logger.warning("EMAIL_FLOW user=- email=%s user_found=0 smtp_ok=0 api_ok=0 reason=user_not_found_for_resend", email)
    else:
        logger.warning("EMAIL_FLOW user=- email=- user_found=0 smtp_ok=0 api_ok=0 reason=empty_email_input")
    return redirect(reverse('verify_email_sent') + '?resent=0')


@require_GET
def request_verification_email_page(request):
    return render(request, 'tools/request_verification_email.html')

def calculator(request):
    return render(request, 'tools/calculator.html')

# Single finance calculator (one panel per page)
CALC_PANELS = ('percent', 'iva', 'salary', 'interest', 'currency', 'loan', 'discount', 'margin', 'split')
CALC_TITLES = {
    'percent': _('Percentage'),
    'iva': _('VAT'),
    'salary': _('Net Salary'),
    'interest': _('Interest'),
    'currency': _('Currency'),
    'loan': _('Loan'),
    'discount': _('Discount'),
    'margin': _('Margin'),
    'split': _('Split Bill'),
}

def calculator_single(request, panel):
    if panel not in CALC_PANELS:
        from django.http import Http404
        raise Http404
    return render(request, 'tools/calculator.html', {
        'active_panel': panel,
        'single_calc': True,
        'calc_title': CALC_TITLES.get(panel, _('Finance calculator')),
    })


CURRENCY_CODES = frozenset({'EUR', 'USD', 'GBP', 'CHF', 'JPY'})


@require_GET
def currency_convert_api(request):
    """Proxy tassi di cambio: richieste same-origin dal browser, niente CORS."""
    try:
        amount = float(request.GET.get('amount', ''))
    except (TypeError, ValueError):
        return JsonResponse({'error': 'invalid_amount'}, status=400)
    if amount < 0:
        return JsonResponse({'error': 'invalid_amount'}, status=400)

    from_curr = (request.GET.get('from') or '').upper().strip()
    to_curr = (request.GET.get('to') or '').upper().strip()
    if from_curr not in CURRENCY_CODES or to_curr not in CURRENCY_CODES:
        return JsonResponse({'error': 'invalid_currency'}, status=400)
    if from_curr == to_curr:
        return JsonResponse({'error': 'same_currency'}, status=400)

    ua = {'User-Agent': 'Mozilla/5.0 (compatible; GadlyTools/1.0)'}
    ff_url = (
        'https://api.frankfurter.app/latest?'
        f'amount={quote(str(amount))}&from={from_curr}&to={to_curr}'
    )
    try:
        req = Request(ff_url, headers=ua)
        with urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode('utf-8'))
        converted = (data.get('rates') or {}).get(to_curr)
        if converted is not None:
            return JsonResponse({
                'amount': amount,
                'from': from_curr,
                'to': to_curr,
                'converted': float(converted),
                'date': data.get('date') or '',
            })
    except (URLError, HTTPError, json.JSONDecodeError, ValueError, socket.timeout, OSError) as exc:
        logger.warning('currency_convert_api: frankfurter failed: %s', exc)

    fb_url = f'https://open.er-api.com/v6/latest/{from_curr}'
    try:
        req = Request(fb_url, headers=ua)
        with urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode('utf-8'))
        rate = (data.get('rates') or {}).get(to_curr)
        if rate is None:
            raise ValueError('no_rate')
        return JsonResponse({
            'amount': amount,
            'from': from_curr,
            'to': to_curr,
            'converted': amount * float(rate),
            'date': '',
        })
    except (URLError, HTTPError, json.JSONDecodeError, ValueError, socket.timeout, OSError) as exc:
        logger.warning('currency_convert_api: fallback failed: %s', exc)
        return JsonResponse({'error': 'upstream'}, status=502)

def translator(request):
    return render(request, 'tools/translator.html')

def file_analyzer(request):
    return render(request, 'tools/file_analyzer.html')

def image_editor(request):
    return render(request, 'tools/image_editor.html')

def data_viz(request):
    return render(request, 'tools/data_viz.html')


def qr_generator(request):
    return render(request, 'tools/qr_generator.html')


def qr_decoder(request):
    return render(request, 'tools/qr_decoder.html')


def password_gen(request):
    return render(request, 'tools/password_gen.html')


def help_password_gen(request):
    return render(request, 'tools/help_password_gen.html')


def remove_spaces(request):
    return render(request, 'tools/remove_spaces.html')

def username_generator(request):
    return render(request, 'tools/username_generator.html')

def bio_generator(request):
    return render(request, 'tools/bio_generator.html')

def markdown_preview(request):
    return render(request, 'tools/markdown_preview.html')


def diff_checker(request):
    return render(request, 'tools/diff_checker.html')


def json_formatter(request):
    return render(request, 'tools/json_formatter.html')


def base64_encoder(request):
    return render(request, 'tools/base64_encoder.html')


def case_converter(request):
    return render(request, 'tools/case_converter.html')


def uuid_generator(request):
    return render(request, 'tools/uuid_generator.html')


def lorem_ipsum(request):
    return render(request, 'tools/lorem_ipsum.html')


def regex_tester(request):
    return render(request, 'tools/regex_tester.html')


def cron_explainer(request):
    return render(request, 'tools/cron_explainer.html')


def hash_generator(request):
    return render(request, 'tools/hash_generator.html')


def color_converter(request):
    return render(request, 'tools/color_converter.html')


def meta_tag_checker(request):
    return render(request, 'tools/meta_tag_checker.html')

def sitemap_extractor(request):
    return render(request, 'tools/sitemap_extractor.html')

def robots_txt_generator(request):
    return render(request, 'tools/robots_txt_generator.html')

def site_speed_check(request):
    return render(request, 'tools/site_speed_check.html')

def hashtag_generator(request):
    return render(request, 'tools/hashtag_generator.html')

def caption_generator(request):
    return render(request, 'tools/caption_generator.html')

def cv_generator(request):
    return render(request, 'tools/cv_generator.html')

def cv_preview_web(request):
    return render(request, 'tools/cv_preview_web.html')

def cv_optimizer(request):
    return render(request, 'tools/cv_optimizer.html')

def cover_letter_generator(request):
    return render(request, 'tools/cover_letter_generator.html')

def application_email_generator(request):
    return render(request, 'tools/application_email_generator.html')

def interview_simulator(request):
    return render(
        request,
        "tools/interview_simulator.html",
        {
            "interview_payload": get_interview_simulator_payload(),
            "interview_ai_config": {
                "enabled": bool(getattr(settings, "INTERVIEW_AI_ENABLED", False)),
                "endpoint": reverse("interview_simulator_ai_reply"),
                "timeoutMs": 7000,
            },
        },
    )

def text_tools_page(request):
    return render(request, 'tools/text_tools_page.html')

def social_media_page(request):
    return render(request, 'tools/social_media_page.html')

def career_cv_page(request):
    return render(request, 'tools/career_cv_page.html')

def file_image_page(request):
    return render(request, 'tools/file_image_page.html')

def finance_page(request):
    return render(request, 'tools/finance_page.html')

def web_seo_page(request):
    return render(request, 'tools/web_seo_page.html')

def security_misc_page(request):
    return render(request, 'tools/security_misc_page.html')

def top_tools_page(request):
    return render(request, 'tools/top_tools_page.html')


# -------------------------
# Help pages for tools
# -------------------------
def help_robots_txt(request):
    return render(request, 'tools/help_robots_txt.html')


def help_json_guide(request):
    return render(request, 'tools/help_json_guide.html')


def help_data_viz(request):
    return render(request, 'tools/help_data_viz.html')


def help_diff_checker(request):
    return render(request, 'tools/help_diff_checker.html')


def help_text_tools(request):
    return render(request, 'tools/help_text_tools.html')


def help_cv_generator(request):
    return render(request, 'tools/help_cv_generator.html')


def help_cv_optimizer(request):
    return render(request, 'tools/help_cv_optimizer.html')


def help_cover_letter(request):
    return render(request, 'tools/help_cover_letter.html')


def help_application_email(request):
    return render(request, 'tools/help_application_email.html')


def help_interview_simulator(request):
    return render(request, 'tools/help_interview_simulator.html')


def help_cron_explainer(request):
    return render(request, 'tools/help_cron_explainer.html')


@require_http_methods(["POST"])
def interview_simulator_ai_reply(request):
    if not getattr(settings, "INTERVIEW_AI_ENABLED", False):
        return JsonResponse({"error": "disabled"}, status=503)

    try:
        data = json.loads(request.body or "{}")
    except Exception:
        return JsonResponse({"error": "invalid_json"}, status=400)

    intent = str(data.get("intent") or "").strip().lower()
    question = str(data.get("question") or "").strip()
    answer = str(data.get("answer") or "").strip()
    lang = str(data.get("lang") or "en").strip().lower()
    recruiter_gender = str(data.get("recruiter_gender") or "male").strip().lower()
    transcript_tail = data.get("transcript_tail") or []

    if intent not in {"feedback", "reply_candidate_question", "encouragement", "clarify"}:
        return JsonResponse({"error": "invalid_intent"}, status=400)

    tail_text = ""
    if isinstance(transcript_tail, list):
        tail_text = "\n".join(str(x) for x in transcript_tail[-6:])

    is_it = lang.startswith("it")
    sys_prompt = (
        "You are a professional job interviewer. "
        "Write concise recruiter messages (max 2 short sentences), practical and natural. "
        "Do not use markdown. Do not mention being an AI."
    )
    if is_it:
        sys_prompt += " Rispondi in italiano naturale."
    else:
        sys_prompt += " Reply in natural English."

    user_prompt = (
        f"Intent: {intent}\n"
        f"Recruiter gender: {recruiter_gender}\n"
        f"Current question: {question}\n"
        f"Candidate answer/message: {answer}\n"
        f"Recent transcript:\n{tail_text}\n"
        "Output only the recruiter message text."
    )

    endpoint = getattr(settings, "INTERVIEW_AI_ENDPOINT", "").strip()
    model = getattr(settings, "INTERVIEW_AI_MODEL", "").strip() or "llama3.1:8b"
    timeout_seconds = int(getattr(settings, "INTERVIEW_AI_TIMEOUT_SECONDS", 8))
    if not endpoint:
        return JsonResponse({"error": "endpoint_missing"}, status=500)

    payload = {
        "model": model,
        "prompt": sys_prompt + "\n\n" + user_prompt,
        "stream": False,
        "options": {"temperature": 0.5, "num_predict": 120},
    }

    req = Request(
        endpoint,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urlopen(req, timeout=timeout_seconds) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
        parsed = json.loads(raw or "{}")
        text = str(parsed.get("response") or "").strip()
        if not text:
            return JsonResponse({"error": "empty_reply"}, status=502)
        if len(text) > 280:
            text = text[:280].rstrip()
        return JsonResponse({"reply": text, "source": "local_ai"})
    except (HTTPError, URLError, socket.timeout):
        return JsonResponse({"error": "upstream_unavailable"}, status=502)
    except Exception:
        return JsonResponse({"error": "upstream_error"}, status=502)

# -------------------------
# Word Counter view
# -------------------------
def word_counter(request):
    count_words = None
    count_chars = None
    text = ''

    if request.method == 'POST':
        text = request.POST.get('text', '')
        if 'count_words_btn' in request.POST:
            count_words = len(text.split())
        elif 'count_chars_btn' in request.POST:
            count_chars = len(text)
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse(
                {
                    'ok': True,
                    'count_words': count_words,
                    'count_chars': count_chars,
                    'text': text,
                }
            )

    return render(request, 'tools/word_counter.html', {
        'count_words': count_words,
        'count_chars': count_chars,
        'text': text
    })

# -------------------------
# API-like view for Translator (server-side)
# -------------------------
@csrf_exempt
def translate_text(request):
    if request.method != "POST":
        return JsonResponse({"error": gettext("Invalid request method")}, status=405)

    try:
        data = json.loads(request.body)
        text = data.get("q", "")  # Il testo da tradurre
        target_lang = data.get("target", "en")  # Lingua di destinazione, di default 'en'

        if not text:
            return JsonResponse({"error": gettext("No text provided")}, status=400)

        # googletrans 4.x usa translate() async: va eseguita con asyncio.run()
        async def _translate():
            async with Translator() as translator:
                return await translator.translate(text, src='auto', dest=target_lang)

        translated = asyncio.run(_translate())
        return JsonResponse({"translatedText": translated.text})

    except Exception as e:
        print("Error in translate_text:", e)
        return JsonResponse({"error": str(e)}, status=500)
# -------------------------
# File Analyzer view
# -------------------------
def _human_size(size_bytes):
    for unit in ('B', 'KB', 'MB', 'GB', 'TB'):
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} PB"

@csrf_exempt
def analyze_file(request):
    if request.method == "POST" and request.FILES.get("file"):
        file = request.FILES["file"]

        file_name = file.name
        file_size = file.size
        file_type = file.content_type or "application/octet-stream"
        file_extension = os.path.splitext(file_name)[1] or "-"
        file_size_human = _human_size(file_size)

        # Hashes
        file.seek(0)
        md5_hasher = hashlib.md5()
        sha256_hasher = hashlib.sha256()
        for chunk in file.chunks():
            md5_hasher.update(chunk)
            sha256_hasher.update(chunk)
        md5_hash = md5_hasher.hexdigest()
        sha256_hash = sha256_hasher.hexdigest()

        # Last modified (optional, provided by browser file metadata).
        last_modified_raw = (request.POST.get("last_modified") or "").strip()
        if last_modified_raw.isdigit():
            try:
                ts = int(last_modified_raw) / 1000.0
                last_modified = time.strftime("%Y-%m-%d %H:%M", time.localtime(ts))
            except Exception:
                last_modified = "-"
        else:
            last_modified = "-"

        result = {
            "file_name": file_name,
            "file_size": file_size,
            "file_size_human": file_size_human,
            "file_type": file_type,
            "file_extension": file_extension,
            "md5_hash": md5_hash,
            "sha256_hash": sha256_hash,
            "last_modified": last_modified,
        }

        # Prova sempre a leggere come testo; se riesce, aggiungi linee/parole/caratteri
        try:
            file.seek(0)
            content = file.read().decode("utf-8", errors="ignore")
            lines = content.splitlines()
            words = content.split()
            result["line_count"] = len(lines)
            result["word_count"] = len(words)
            result["char_count"] = len(content)
            result["char_count_no_spaces"] = len(re.sub(r"\s+", "", content))
            result["reading_time_min"] = max(1, int((len(words) + 199) / 200)) if words else 0
        except Exception:
            pass

        return JsonResponse(result)

    return JsonResponse({"error": gettext("No file uploaded")}, status=400)

# -----------------------------------------
# PDF Merger view
# -----------------------------------------
@csrf_exempt
def pdf_merger(request):
    if request.method == 'POST' and request.FILES.getlist('files'):
        files = request.FILES.getlist('files')
        
        writer = PdfWriter()

        # Unisci ogni PDF caricato
        for file in files:
            reader = PdfReader(file)
            for page in reader.pages:
                writer.add_page(page)

        # Usa tempfile per creare un file temporaneo per il PDF unito
        try:
            temp_dir = os.path.join(tempfile.gettempdir(), 'gadly_pdf_merge')
            os.makedirs(temp_dir, exist_ok=True)
            fd, output_path = tempfile.mkstemp(suffix='.pdf', dir=temp_dir)
            os.close(fd)

            # Scrivi il PDF unito nel file temporaneo
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)

            # Restituisci il PDF come file da scaricare
            with open(output_path, 'rb') as output_file:
                pdf_data = output_file.read()
            try:
                os.unlink(output_path)
            except Exception:
                pass
            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="merged_pdf.pdf"'
            return response

        except Exception as e:
            return JsonResponse({"message": gettext("Error: %(err)s") % {'err': str(e)}}, status=500)

    return render(request, 'tools/pdf_merge.html')


# -----------------------------------------
# PDF Split view
# -----------------------------------------
@csrf_exempt
def pdf_split(request):
    if request.method == 'POST' and request.FILES.get('file'):
        pdf_file = request.FILES['file']
        pages_per_split = 1
        if request.POST.get('pages_per_split'):
            try:
                pages_per_split = max(1, int(request.POST['pages_per_split']))
            except (ValueError, TypeError):
                pages_per_split = 1
        zip_path = None
        try:
            reader = PdfReader(pdf_file)
            num_pages = len(reader.pages)
            if num_pages == 0:
                return JsonResponse({"error": gettext("PDF has no pages")}, status=400)
            base_name = os.path.splitext(pdf_file.name)[0]
            fd, zip_path = tempfile.mkstemp(suffix='.zip')
            os.close(fd)
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                chunk_num = 0
                i = 0
                while i < num_pages:
                    writer = PdfWriter()
                    end_i = min(i + pages_per_split, num_pages)
                    for j in range(i, end_i):
                        writer.add_page(reader.pages[j])
                    chunk_num += 1
                    chunk_name = f"{base_name}_part_{chunk_num}.pdf"
                    fd2, tmp_path = tempfile.mkstemp(suffix='.pdf')
                    os.close(fd2)
                    try:
                        with open(tmp_path, 'wb') as out:
                            writer.write(out)
                        zf.write(tmp_path, chunk_name)
                    finally:
                        if os.path.exists(tmp_path):
                            os.unlink(tmp_path)
                    i = end_i
            with open(zip_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/zip')
                response['Content-Disposition'] = f'attachment; filename="{base_name}_split.zip"'
                return response
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
        finally:
            if zip_path and os.path.exists(zip_path):
                try:
                    os.unlink(zip_path)
                except Exception:
                    pass
    return render(request, 'tools/pdf_split.html')


# -----------------------------------------
# Image Editor API: Background removal & Upscaling
# -----------------------------------------
import io

def _get_rembg():
    """Lazy import rembg to avoid blocking app startup if not installed."""
    try:
        from rembg import remove as rembg_remove
        return rembg_remove
    except Exception:
        return None


@csrf_exempt
def image_remove_bg_api(request):
    """Remove background from image. Requires: pip install "rembg[cpu]" """
    if request.method != 'POST' or not request.FILES.get('image'):
        return JsonResponse({"error": gettext("POST with image file required")}, status=400)
    rembg_remove = _get_rembg()
    if rembg_remove is None:
        return JsonResponse({"error": gettext('Background removal requires: pip install "rembg[cpu]"')}, status=501)
    try:
        from PIL import Image as PILImage
        request.FILES['image'].seek(0)
        img = PILImage.open(request.FILES['image']).copy().convert('RGBA')
        w, h = img.size
        max_side = 4096
        if w > max_side or h > max_side:
            r = PILImage.Resampling.LANCZOS if hasattr(PILImage, 'Resampling') else PILImage.LANCZOS
            img.thumbnail((max_side, max_side), r)
            w, h = img.size
        out = rembg_remove(
            img,
            post_process_mask=False,
            alpha_matting=True,
            alpha_matting_foreground_threshold=285,
            alpha_matting_background_threshold=15,
        )
        buf = io.BytesIO()
        out.save(buf, format='PNG')
        buf.seek(0)
        return HttpResponse(buf.getvalue(), content_type='image/png')
    except Exception as e:
        err_msg = str(e)
        if 'onnxruntime' in err_msg.lower() or 'session' in err_msg.lower():
            err_msg = gettext('rembg/onnxruntime error. Run: pip install "rembg[cpu]"')
        elif not err_msg or len(err_msg) > 200:
            err_msg = err_msg[:200] if err_msg else gettext("Unknown error")
        return JsonResponse({"error": err_msg}, status=500)


@csrf_exempt
def image_upscale_api(request):
    """Upscale image 2x using high-quality resampling."""
    if request.method != 'POST' or not request.FILES.get('image'):
        return JsonResponse({"error": gettext("POST with image file required")}, status=400)
    try:
        from PIL import Image as PILImage
        f = request.FILES['image']
        f.seek(0)
        scale = 2
        if request.POST.get('scale'):
            try:
                scale = max(1, min(4, int(request.POST['scale'])))
            except (ValueError, TypeError):
                pass
        img = PILImage.open(f).copy().convert('RGBA')
        w, h = img.size
        new_w, new_h = w * scale, h * scale
        if new_w > 4096 or new_h > 4096:
            return JsonResponse({"error": gettext("Result would exceed 4096px. Use smaller image or scale.")}, status=400)
        try:
            resample = PILImage.Resampling.LANCZOS
        except AttributeError:
            resample = PILImage.LANCZOS
        out = img.resize((new_w, new_h), resample)
        buf = io.BytesIO()
        out.save(buf, format='PNG')
        return HttpResponse(buf.getvalue(), content_type='image/png')
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


# -----------------------------------------
# Web / SEO API views
# -----------------------------------------
def _fetch_url(url, timeout=15):
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    req = Request(url, headers={'User-Agent': 'Mozilla/5.0 (compatible; ToolsBot/1.0)'})
    return urlopen(req, timeout=timeout)


@csrf_exempt
def meta_tag_check_api(request):
    with override(_request_language(request)):
        if request.method != 'POST':
            return JsonResponse({"error": gettext("Invalid method")}, status=405)
        try:
            data = json.loads(request.body)
            url = data.get('url', '').strip()
            if not url:
                return JsonResponse({"error": gettext("URL required")}, status=400)
            resp = _fetch_url(url)
            html = resp.read().decode('utf-8', errors='ignore')
            meta_tags = []
            # title
            m = re.search(r'<title[^>]*>([^<]+)</title>', html, re.I | re.S)
            if m:
                meta_tags.append({"name": "title", "content": m.group(1).strip()})
            # meta tags
            for m in re.finditer(r'<meta\s+([^>]+)>', html, re.I):
                attrs = m.group(1)
                name = None
                prop = None
                content = None
                for a in re.finditer(r'(\w+)\s*=\s*["\']([^"\']*)["\']', attrs):
                    k, v = a.group(1).lower(), a.group(2)
                    if k == 'name': name = v
                    elif k == 'property': prop = v
                    elif k == 'content': content = v
                if content and (name or prop):
                    meta_tags.append({"name": name or prop, "content": content})
            return JsonResponse({"meta_tags": meta_tags})
        except (URLError, HTTPError):
            return JsonResponse(
                {"error": gettext("The URL could not be reached. Check the address and try again.")},
                status=400,
            )
        except Exception:
            return JsonResponse(
                {"error": gettext("An error occurred while checking meta tags.")},
                status=500,
            )


@csrf_exempt
def sitemap_extract_api(request):
    if request.method != 'POST':
        return JsonResponse({"error": gettext("Invalid method")}, status=405)
    try:
        data = json.loads(request.body)
        url = data.get('url', '').strip()
        if not url:
            return JsonResponse({"error": gettext("URL required")}, status=400)
        resp = _fetch_url(url)
        content = resp.read()
        # Check if we got HTML instead of XML
        try:
            peek = content[:100].decode('utf-8', errors='ignore').strip().lower()
        except Exception:
            peek = ''
        if peek.startswith('<!') or '<html' in peek[:50]:
            return JsonResponse({
                "error": gettext(
                    "This URL returns a web page (HTML), not a sitemap. Please enter the direct link to your sitemap file (e.g. https://example.com/sitemap.xml)."
                )
            }, status=400)
        root = ET.fromstring(content)
        ns = {'sm': 'http://www.sitemaps.org/schemas/sitemap/0.9'}
        urls = []
        for loc in root.findall('.//sm:url/sm:loc', ns):
            if loc.text:
                urls.append(loc.text.strip())
        if not urls:
            for loc in root.findall('.//{http://www.sitemaps.org/schemas/sitemap/0.9}url/{http://www.sitemaps.org/schemas/sitemap/0.9}loc'):
                if loc.text:
                    urls.append(loc.text.strip())
        if not urls:
            for loc in root.findall('.//*[local-name()="loc"]'):
                if loc.text:
                    urls.append(loc.text.strip())
        return JsonResponse({"urls": urls})
    except (URLError, HTTPError):
        return JsonResponse({
            "error": gettext(
                "The URL could not be reached. Check the address and try again."
            )
        }, status=400)
    except ET.ParseError as e:
        return JsonResponse({
            "error": gettext(
                "This URL does not return valid sitemap XML. Please enter the direct link to your sitemap file (e.g. https://example.com/sitemap.xml)."
            )
        }, status=400)
    except Exception:
        return JsonResponse({
            "error": gettext(
                "An error occurred while extracting sitemap URLs."
            )
        }, status=500)


def _robots_txt_url(url):
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    parsed = urlparse(url)
    path = (parsed.path or '').rstrip('/')
    if path.endswith('robots.txt'):
        robots_path = path
    else:
        robots_path = '/robots.txt'
    return urlunparse((parsed.scheme, parsed.netloc, robots_path, '', '', ''))


def _parse_robots_txt(text):
    """Return fields for the * user-agent block, or the first block if none."""
    lines = []
    for raw in text.splitlines():
        line = raw.split('#', 1)[0].strip()
        if line:
            lines.append(line)
    blocks = []
    current = {'agents': [], 'disallow': [], 'allow': []}
    sitemaps = []

    def flush():
        nonlocal current
        if current['agents'] or current['disallow'] or current['allow']:
            blocks.append(current)
        current = {'agents': [], 'disallow': [], 'allow': []}

    for line in lines:
        if ':' not in line:
            continue
        key, _, val = line.partition(':')
        key = key.strip().lower()
        val = val.strip()
        if key == 'user-agent':
            if current['agents'] or current['disallow'] or current['allow']:
                flush()
            current['agents'].append(val)
        elif key == 'disallow':
            current['disallow'].append(val)
        elif key == 'allow':
            current['allow'].append(val)
        elif key == 'sitemap' and val:
            sitemaps.append(val)
    flush()

    chosen = None
    for block in blocks:
        if '*' in block['agents']:
            chosen = block
            break
    if not chosen and blocks:
        chosen = blocks[0]

    user_agent = '*'
    disallow = []
    allow = []
    if chosen:
        if chosen['agents']:
            user_agent = chosen['agents'][0]
        disallow = [p for p in chosen['disallow'] if p]
        allow = [p for p in chosen['allow'] if p]

    return {
        'user_agent': user_agent,
        'disallow': disallow,
        'allow': allow,
        'sitemap': sitemaps[0] if sitemaps else '',
        'sitemaps': sitemaps,
    }


@csrf_exempt
def robots_fetch_api(request):
    with override(_request_language(request)):
        if request.method != 'POST':
            return JsonResponse({"error": gettext("Invalid method")}, status=405)
        try:
            data = json.loads(request.body)
            url = data.get('url', '').strip()
            if not url:
                return JsonResponse(
                    {"error": gettext("URL required"), "code": "url_required"},
                    status=400,
                )
            robots_url = _robots_txt_url(url)
            resp = _fetch_url(robots_url)
            content = resp.read().decode('utf-8', errors='ignore')
            peek = content[:200].strip().lower()
            if peek.startswith('<!') or '<html' in peek[:80]:
                return JsonResponse({
                    "error": gettext(
                        "No robots.txt found at this address (the server returned a web page). Check the site URL."
                    ),
                    "code": "no_robots",
                }, status=400)
            parsed = _parse_robots_txt(content)
            return JsonResponse({
                "content": content,
                "robots_url": robots_url,
                "parsed": parsed,
            })
        except (URLError, HTTPError):
            return JsonResponse({
                "error": gettext(
                    "The URL could not be reached. Check the address and try again."
                ),
                "code": "unreachable",
            }, status=400)
        except Exception:
            return JsonResponse({
                "error": gettext(
                    "An error occurred while loading robots.txt."
                ),
                "code": "fetch_error",
            }, status=500)


@require_GET
def vat_rates_api(request):
    """Proxy verso vat-api.eu: richieste same-origin dal browser, niente CORS."""
    upstream = "https://vat-api.eu/api/v1/rates"
    try:
        req = Request(
            upstream,
            headers={"User-Agent": "Mozilla/5.0 (compatible; GadlyTools/1.0; +https://vat-api.eu)"},
        )
        with urlopen(req, timeout=15) as resp:
            body = resp.read().decode("utf-8")
        data = json.loads(body)
        if not isinstance(data, list):
            return JsonResponse({"error": "unexpected_shape"}, status=502)
        return JsonResponse(data, safe=False)
    except (URLError, HTTPError, json.JSONDecodeError, ValueError, socket.timeout, OSError) as exc:
        logger.warning("vat_rates_api: upstream failed: %s", exc)
        return JsonResponse({"error": "upstream"}, status=502)


@csrf_exempt
def site_speed_api(request):
    with override(_request_language(request)):
        if request.method != 'POST':
            return JsonResponse({"error": gettext("Invalid method")}, status=405)
        try:
            data = json.loads(request.body)
            url = data.get('url', '').strip()
            if not url:
                return JsonResponse({"error": gettext("URL required")}, status=400)
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            req = Request(url, headers={'User-Agent': 'Mozilla/5.0 (compatible; ToolsBot/1.0)'})
            t0 = time.time()
            resp = urlopen(req, timeout=30)
            _ = resp.read()
            elapsed = round((time.time() - t0) * 1000)
            return JsonResponse({"time_ms": elapsed, "status": resp.status})
        except (URLError, HTTPError) as e:
            if isinstance(e, URLError) and isinstance(getattr(e, "reason", None), socket.timeout):
                return JsonResponse({"error": gettext("The request timed out. Please try again.")}, status=408)
            return JsonResponse({
                "error": gettext("The URL could not be reached. Check the address and try again.")
            }, status=400)
        except socket.timeout:
            return JsonResponse({"error": gettext("The request timed out. Please try again.")}, status=408)
        except Exception as e:
            if "timed out" in str(e).lower():
                return JsonResponse({"error": gettext("The request timed out. Please try again.")}, status=408)
            return JsonResponse({
                "error": gettext("The URL could not be reached. Check the address and try again.")
            }, status=500)


def _extract_text_from_uploaded_cv(uploaded_file):
    ext = os.path.splitext(uploaded_file.name or "")[1].lower()
    uploaded_file.seek(0)

    if ext == ".txt":
        return uploaded_file.read().decode("utf-8", errors="ignore")

    if ext == ".pdf":
        reader = PdfReader(uploaded_file)
        chunks = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if not page_text and "/Rotate" in page:
                try:
                    page_text = (page.extract_text() or "")
                except Exception:
                    page_text = ""
            chunks.append(page_text)
        return "\n".join(chunks)

    if ext == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise RuntimeError(gettext("DOCX support is not available on the server.")) from exc
        doc = Document(uploaded_file)
        text_chunks = [p.text for p in doc.paragraphs if p.text]
        try:
            for tbl in doc.tables:
                for row in tbl.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if row_cells:
                        text_chunks.append(" | ".join(row_cells))
        except Exception:
            pass
        return "\n".join(text_chunks)

    raise ValueError(gettext("Unsupported file type. Please upload PDF, DOCX, or TXT."))


def _tokenize_keywords(text):
    raw = re.findall(r"[a-zA-Z][a-zA-Z0-9\\-\\+]{2,}", (text or "").lower())
    stop_words = {
        "the", "and", "for", "with", "from", "that", "this", "your", "you",
        "per", "con", "dei", "delle", "della", "dello", "sulla", "sono", "una", "uno",
        "job", "role", "position", "about"
    }
    return [w for w in raw if w not in stop_words]


def _sanitize_extracted_text(text):
    cleaned = (text or "").replace("\x00", " ")
    cleaned = re.sub(r"[\u2022\u25CF\u25A0\u25AA]", "- ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned, flags=re.UNICODE)
    return cleaned.strip()


def _extract_contact_fields(cv_text):
    text = cv_text or ""
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(?:\+?\d[\d\s().-]{6,}\d)", text)
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s,;]+", text, flags=re.IGNORECASE)
    return {
        "email": email_match.group(0).strip() if email_match else "",
        "phone": phone_match.group(0).strip() if phone_match else "",
        "linkedin": linkedin_match.group(0).strip() if linkedin_match else "",
    }


def cv_optimize_api(request):
    if request.method != "POST":
        return JsonResponse({"error": gettext("Invalid method")}, status=405)

    try:
        uploaded = request.FILES.get("cv_file")
        if not uploaded:
            return JsonResponse({"error": gettext("Please upload your CV file.")}, status=400)

        if uploaded.size > 8 * 1024 * 1024:
            return JsonResponse({"error": gettext("File too large. Maximum allowed size is 8 MB.")}, status=400)

        job_title = (request.POST.get("job_title") or "").strip()
        job_description = (request.POST.get("job_description") or "").strip()

        cv_text = _sanitize_extracted_text(_extract_text_from_uploaded_cv(uploaded))
        if not cv_text:
            return JsonResponse({"error": gettext("Could not extract readable text from this file.")}, status=400)
        extracted_contact = _extract_contact_fields(cv_text)

        lower_text = cv_text.lower()
        section_signals = {
            "Summary": ["summary", "profile", "objective", "about me", "profilo", "obiettivo"],
            "Experience": ["experience", "work", "employment", "esperienza", "lavoro"],
            "Education": ["education", "degree", "university", "formazione", "istruzione", "laurea"],
            "Skills": ["skills", "competencies", "competenze", "abilità", "abilita"],
            "Contact": ["email", "phone", "telefono", "linkedin", "github"],
        }
        missing_sections = []
        for section_name, signals in section_signals.items():
            if not any(sig in lower_text for sig in signals):
                missing_sections.append(section_name)

        words = re.findall(r"\\w+", cv_text, flags=re.UNICODE)
        word_count = len(words)
        score = 100
        suggestions = []

        if word_count < 180:
            score -= 20
            suggestions.append(gettext("Your CV is very short. Add more detail about impact and achievements."))
        elif word_count > 950:
            score -= 12
            suggestions.append(gettext("Your CV is quite long. Try to keep it concise and focused."))

        if missing_sections:
            score -= min(30, len(missing_sections) * 7)
            suggestions.append(
                gettext("Add or improve these sections: %(sections)s")
                % {"sections": ", ".join(missing_sections)}
            )

        keyword_source = " ".join([job_title, job_description]).strip()
        jd_keywords = _tokenize_keywords(keyword_source)
        matched_keywords = []
        missing_keywords = []
        if jd_keywords:
            seen = set()
            for kw in jd_keywords:
                if kw in seen:
                    continue
                seen.add(kw)
                if kw in lower_text:
                    matched_keywords.append(kw)
                else:
                    missing_keywords.append(kw)
            match_ratio = (len(matched_keywords) / len(seen)) if seen else 0
            if match_ratio < 0.35:
                score -= 20
                suggestions.append(gettext("Low keyword match with the target role. Tailor your CV to the job description."))
            elif match_ratio < 0.60:
                score -= 10
                suggestions.append(gettext("Good start, but you can include more role-specific keywords."))
        else:
            match_ratio = None

        score = max(1, min(100, score))
        if not suggestions:
            suggestions.append(gettext("Great baseline CV. Fine-tune wording and quantifiable results for stronger impact."))

        return JsonResponse({
            "score": score,
            "word_count": word_count,
            "missing_sections": missing_sections,
            "matched_keywords": matched_keywords[:15],
            "missing_keywords": missing_keywords[:15],
            "suggestions": suggestions[:6],
            "extracted_contact": extracted_contact,
            "keyword_match_ratio": match_ratio,
        })
    except ValueError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except RuntimeError as exc:
        return JsonResponse({"error": str(exc)}, status=500)
    except Exception:
        return JsonResponse({"error": gettext("An error occurred while optimizing the CV.")}, status=500)


@require_GET
def robots_txt(request):
    """Root robots.txt: disallow admin/accounts/api; point crawlers to sitemap.xml."""
    sitemap_abs = request.build_absolute_uri(reverse("sitemap_xml"))
    lines = [
        "User-agent: *",
        "Disallow: /admin/",
        "Disallow: /accounts/",
        "Disallow: /api/",
        "Disallow: /jsi18n/",
        "",
        f"Sitemap: {sitemap_abs}",
        "",
    ]
    return HttpResponse("\n".join(lines), content_type="text/plain; charset=utf-8")

